from pathlib import Path
import base64, re, shutil, subprocess, zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('docs/sq'); OUT=Path('word-pack'); LOGO=Path('/tmp/ltline-logo.png')
LOGO_B64=Path('assets/ltline-logo.png.b64').read_text(encoding='utf-8').strip()
BUILD_DATE='30.08.2026'

def meta(md):
    def g(key, default):
        pats=[rf'\*\*{re.escape(key)}\s*:\s*\*\*\s*([^|\n]+)', rf'\*\*{re.escape(key)}\*\*\s*:\s*([^|\n]+)']
        for pat in pats:
            m=re.search(pat,md,re.I)
            if m: return m.group(1).strip()
        return default
    m=re.search(r'^#\s+(.+)$',md,re.M); title=m.group(1).strip() if m else 'LTLINE'
    title=re.sub(r'^LTLINE\s*[—-]\s*','',title).strip()
    return title,g('ID e dokumentit',g('ID','TBD')),g('Revizioni','1.0'),g('Statusi','PROJEKT')

def borderless(table):
    pr=table._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement('w:'+e); x.set(qn('w:val'),'nil'); b.append(x)
    pr.append(b)

def field(p,code):
    f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),code); p._p.append(f)

def style_doc(doc):
    for name,size,bold in [('Normal',10.5,False),('Title',22,True),('Heading 1',15,True),('Heading 2',12.5,True),('Heading 3',11,True)]:
        st=doc.styles[name]; st.font.name='Aptos'; st.font.size=Pt(size); st.font.bold=bold
    s=doc.sections[0]; s.top_margin=Cm(2.8); s.bottom_margin=Cm(2.1); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2); s.header_distance=Cm(.7); s.footer_distance=Cm(.8)

def add_branding(doc,title,doc_id,rev,status):
    sec=doc.sections[0]
    h=sec.header; t=h.add_table(rows=1,cols=2,width=Cm(16.6)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; borderless(t)
    l,r=t.rows[0].cells; l.vertical_alignment=r.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=l.paragraphs[0]; rr=p.add_run(); rr.add_picture(str(LOGO),width=Cm(2.4))
    p=r.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rr=p.add_run('LTLINE\n'); rr.bold=True; rr.font.size=Pt(12)
    rr=p.add_run(title.upper()+'\n'); rr.bold=True; rr.font.size=Pt(8.5)
    rr=p.add_run(f'ID: {doc_id}  |  Revizioni: {rev}  |  Statusi: {status}'); rr.font.size=Pt(8)
    p=h.add_paragraph(); pr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8'); bot.set(qn('w:space'),'1'); pb.append(bot); pr.append(pb)
    f=sec.footer; p=f.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(f'LTLINE — Dokument i Kontrolluar  |  Revizioni {rev}  |  Faqja '); rr.font.size=Pt(8); field(p,'PAGE')
    rr=p.add_run(' nga '); rr.font.size=Pt(8); field(p,'NUMPAGES'); rr=p.add_run('  |  Dokumentacioni zyrtar LTLINE'); rr.font.size=Pt(8)

def insert_control_block_first(doc,title,doc_id,rev,status):
    body=doc._element.body
    anchor=body.find(qn('w:sectPr'))
    nodes=[]
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('LTLINE').bold=True; nodes.append(p._p)
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title.upper()); nodes.append(p._p)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(f'{doc_id}  •  Revizioni {rev}  •  {status}'); rr.bold=True; rr.font.size=Pt(10); nodes.append(p._p)
    ct=doc.add_table(rows=4,cols=4); ct.style='Table Grid'; ct.alignment=WD_TABLE_ALIGNMENT.CENTER
    vals=[('ID e dokumentit',doc_id,'Revizioni',rev),('Statusi',status,'Data',BUILD_DATE),('Përgatiti','TBD','Shqyrtoi','TBD'),('Miratoi','TBD','Data e hyrjes në fuqi','TBD')]
    for i,row in enumerate(vals):
        for j,v in enumerate(row):
            c=ct.cell(i,j); c.text=v; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in c.paragraphs[0].runs: run.font.size=Pt(8.5); run.bold=(j%2==0)
    nodes.append(ct._tbl)
    p=doc.add_paragraph(); nodes.append(p._p)
    for node in nodes:
        body.remove(node)
    for node in nodes: body.insert(0,node)

def build(md_file,out):
    md=md_file.read_text(encoding='utf-8'); title,doc_id,rev,status=meta(md)
    subprocess.run(['pandoc',str(md_file),'-o',str(out),'--from=gfm'],check=True)
    doc=Document(out); style_doc(doc); add_branding(doc,title,doc_id,rev,status); insert_control_block_first(doc,title,doc_id,rev,status); doc.save(out)

def main():
    LOGO.write_bytes(base64.b64decode(LOGO_B64))
    if OUT.exists(): shutil.rmtree(OUT)
    OUT.mkdir(parents=True)
    for md in sorted(ROOT.rglob('*.md')):
        out=(OUT/md.relative_to(ROOT)).with_suffix('.docx'); out.parent.mkdir(parents=True,exist_ok=True); build(md,out)
    (OUT/'README-WORD-PACK.md').write_text('# LTLINE — Word Pack\n\nPaketë Word e gjeneruar nga dokumentacioni shqiptar `docs/sq/`. Çdo dokument përdor standardin LTLINE me logo, header, footer, ID, revizion, status, bllok kontrolli dhe numërim faqesh.\n',encoding='utf-8')
    (OUT/'FILE-LIST.txt').write_text('\n'.join(str(p.relative_to(OUT)) for p in sorted(OUT.rglob('*.docx')))+'\n',encoding='utf-8')
    z=Path('LTLINE-Foundation-Pack-Word-SQ.zip')
    if z.exists(): z.unlink()
    with zipfile.ZipFile(z,'w',zipfile.ZIP_DEFLATED) as a:
        for p in OUT.rglob('*'):
            if p.is_file(): a.write(p,p.relative_to(OUT))
    LOGO.unlink(missing_ok=True)
if __name__=='__main__': main()
